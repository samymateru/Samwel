from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_table_of_content(doc: Document):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Title'
    hdr_cells[2].text = 'Audit Finding Rating'





file = Document()

table = file.add_table(
    rows=1,
    cols=1
)

table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'


hdr_cells = table.rows[0].cells

hdr_cells[0].text = 'No'
print(hdr_cells)

file.save("hello.docx")
