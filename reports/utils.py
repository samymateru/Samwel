import html
import re
import warnings
warnings.filterwarnings("ignore")
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Inches



def set_cell_background_color(cell, color_hex: str):
    """
    Set the background color of a table cell in a Word document.

    Args:
        cell: The cell object from a `docx` table.
        color_hex (str): Hex code like 'FF0000' (red), 'D9D9D9' (gray), etc.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')     # <-- required
    shd.set(qn('w:color'), 'auto')    # <-- required
    shd.set(qn('w:fill'), color_hex)  # <-- actual background color

    # Remove existing <w:shd> if present (to avoid duplicates)
    for child in tcPr.findall(qn('w:shd')):
        tcPr.remove(child)

    tcPr.append(shd)



def set_cell_text_color(cell, color_hex: str):
    """
    Set the text color inside a cell.
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = [RGBColor.from_string(color_hex)]

def color_shading(value: str):
    """Return cell color based on rating/value."""
    rating_map = {
        "Unacceptable": "FF0000",
        "Significant Improvement Required": "FFC000",
        "Improvement Required": "FFFF00",
        "Acceptable": "9250D0"
    }
    return rating_map.get(value) or "FFFFFF"


def create_styled_table(
    doc,
    columns: int,
    data: list[list[str]],
    column_widths: list[float] = None,
    headers: list[str] = None,
    header_bg: str = "00FFC0",
    header_font_color: str = "FFFFFF",
    row_bg: str = "FFFFFF",
    alt_row_bg: str = None,
    row_height: float = 0.3,
    font_name: str = "Arial Narrow",
    font_size: int = 11,
    full_width: bool = True,
):
    """Create a Word table with strict widths and per-cell shading."""

    # --- Create table ---
    table = doc.add_table(rows=1 if headers else 0, cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    # --- Column widths ---
    if full_width:
        total_width = 7  # inches
        if column_widths:
            # scale proportionally
            total_ratio = sum(column_widths)
            column_widths = [total_width * (w / total_ratio) for w in column_widths]
        else:
            column_widths = [total_width / columns] * columns
    elif not column_widths:
        column_widths = [1.5] * columns

    # --- Header ---
    if headers:
        header_row = table.rows[0]
        if row_height:
            header_row.height = Inches(row_height)
            header_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for i, text in enumerate(headers):
            cell = header_row.cells[i]
            cell.width = Inches(column_widths[i])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(sanitize_for_xml(text))
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.name = font_name
            run.font.color.rgb = RGBColor.from_string(header_font_color)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_bg}"/>')
            cell._element.get_or_add_tcPr().append(shading)

    # --- Rows ---
    for r_idx, row_data in enumerate(data):
        row = table.add_row()
        if row_height:
            row.height = Inches(row_height)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = Inches(column_widths[c_idx])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(sanitize_for_xml(str(value)))
            run.font.name = font_name
            run.font.size = Pt(font_size)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # shading
            bg = alt_row_bg if alt_row_bg and r_idx % 2 == 1 else row_bg
            dynamic = color_shading(str(value))
            fill = dynamic if dynamic != "FFFFFF" else bg
            if fill:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
                cell._element.get_or_add_tcPr().append(shading)

            # enforce width using twips
            cell._tc.get_or_add_tcPr().append(
                parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(column_widths[c_idx]*1440)}" w:type="dxa"/>')
            )

    return table



def sanitize_for_xml(text: str) -> str:
    if not text:
        return ""

    text = str(text)
    # Remove control chars
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)
    # Escape unsafe XML entities
    text = re.sub(r"&(?!(?:[a-zA-Z]+|#\d+);)", "&amp;", text)
    text = text.replace("<", "&lt;").replace(">", "&gt;")
    # Remove invalid surrogates
    text = text.encode("utf-8", "ignore").decode("utf-8")
    return html.escape(text)


def add_hyperlink(paragraph, url, text, color="0000EE", underline=True):
    """
    Adds a clickable hyperlink to a paragraph (internal or external).
    """
    part = paragraph.part
    r_id = part.relate_to(str(url), 1, is_external=True)


    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    new_run.append(rPr)
    new_run_text = OxmlElement('w:t')
    new_run_text.text = str(text)
    new_run.append(new_run_text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

#
# doc = Document()
#
#
# headers = ["Name", "Role", "Email"]
#
#
# data = [
#     ["Gerry Trasures", "Audit Lead", "kito@gmail.com"],
#     ["Collin Wilson", "Audit Assistant", "codex6992@gmail.com"],
#     ["Lydia Nyoni", "Reviewer", "lydia@gmail.com"],
# ]
#
# create_styled_table(
#     doc,
#     columns=3,
#     headers=headers,
#     data=data,
#     column_widths=[0.1, 1.5, 2.5],
#     header_bg="000092",          # dark blue header
#     header_font_color="FFFFFF",  # white text
#     row_bg="FFFFFF",             # light blue for data rows
#     alt_row_bg="FFFFFF" ,         # white for alternating rows
#     row_height=0.3
#
# )
#
# doc.save("styled_table.docx")