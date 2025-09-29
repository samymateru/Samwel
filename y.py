from docxtpl import DocxTemplate
from docx import Document

# Your custom render function that parses the Tiptap JSON to a docx.Document
def render_doc(json_data) -> Document:
    doc = Document()
    for block in json_data['content']:
        if block['type'] == 'paragraph':
            para = doc.add_paragraph()
            for span in block.get('content', []):
                run = para.add_run(span['text'])
                # Style detection can go here (bold, italic, etc.)
                if span.get('marks'):
                    for mark in span['marks']:
                        if mark['type'] == 'bold':
                            run.bold = True
                        if mark['type'] == 'italic':
                            run.italic = True
        # Add handling for other block types (e.g. lists) if needed
    return doc

# Step 1: Load your DocxTemplate (this has your fixed content like title, header, etc.)
template = DocxTemplate("name.docx")

# Step 2: Render your template if you want to use placeholders
context = {
    'name': 'My Report',
    # Don't include the parsed content here — we append it manually
}
template.render(context)

# Step 3: Save the rendered base to a temp file
temp_filename = "temp_rendered_template.docx"
template.save(temp_filename)

# Step 4: Load it as a docx.Document
base_doc = Document(temp_filename)

# Step 5: Parse the JSON



json_data = {
    "type": "doc",
    "content": [
        {
            "type": "paragraph",
            "content": [
                { "type": "text", "text": "Hello ", "marks": [
                    {"type": "bold"},
                    {"type": "color", "attrs": {"color": "#FF0000"}},
                    {"type": "underline"},
                    {"type": "fontsize", "attrs": {"size": 36}},
                    {
                      "type": "fontfamily",
                      "attrs": {
                        "name": "Arial"
                      }
                    }
                ] },
                { "type": "text", "text": "world!" }
            ]
        },
        {
            "type": "ordered_list",
            "content": [
                {
                    "type": "list_item",
                    "content": [
                        {"type": "paragraph", "content": [
                            {"type": "text", "text": "Hello ", "marks": [
                                {"type": "bold"},
                                {"type": "color", "attrs": {"color": "#FF0000"}},
                                {"type": "underline"},
                                {"type": "fontsize", "attrs": {"size": 36}},
                                {
                                    "type": "fontfamily",
                                    "attrs": {
                                        "name": "Arial"
                                    }
                                }
                            ]}
                        ]}
                    ]
                },
                {
                    "type": "list_item",
                    "content": [
                        {"type": "paragraph", "content": [{"type": "text", "text": "Second bullet"}]}
                    ]
                }
            ]
        }

    ]
}




parsed_doc = render_doc(json_data)

# Step 6: Append paragraphs from parsed_doc into base_doc
for para in parsed_doc.paragraphs:
    new_para = base_doc.add_paragraph()
    for run in para.runs:
        r = new_para.add_run(run.text)
        r.bold = run.bold
        r.italic = run.italic
        r.underline = run.underline
        r.font.color.rgb = run.font.color.rgb
        r.font.name = run.font.name
        r.style = run.style

# Step 7: Save final combined document
base_doc.save("final_output.docx")
