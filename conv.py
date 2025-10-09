import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Pt
from reports.utils import sanitize_for_xml


def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))




def strip_word_xml(text: str) -> str:
    """
    Remove escaped WordML (e.g., &lt;w:p&gt;&lt;w:t&gt;...&lt;/w:t&gt;&lt;/w:p&gt;)
    and other unwanted markup from text.
    """
    if not text:
        return ""
    # Unescape < and >
    text = text.replace("&lt;", "<").replace("&gt;", ">")

    # Remove actual WordML tags
    text = re.sub(r"<w:[^>]+>", "", text)
    text = re.sub(r"</w:[^>]+>", "", text)

    # Remove any leftover XML-like tags
    text = re.sub(r"<[^>]+>", "", text)

    return text.strip()


def sanitize_for_xml(text: str) -> str:
    """
    Keep only printable characters and clean escaped tags.
    """
    text = strip_word_xml(text)
    return "".join(c for c in text if c.isprintable() or c in "\n\t\r")


# --------------------------------------------------------
# 🎨 Text Formatting Helpers
# --------------------------------------------------------

def hex_to_rgb(hex_color: str):
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))


def apply_marks(run, marks):
    """Apply text formatting marks (bold, color, etc.) to a run."""
    for mark in marks:
        mtype = mark["type"]
        if mtype == "bold":
            run.bold = True
        elif mtype == "italic":
            run.italic = True
        elif mtype == "underline":
            run.underline = True
        elif mtype == "fontfamily":
            font_name = mark.get("attrs", {}).get("name", "")
            if font_name:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        elif mtype == "color":
            color_code = mark.get("attrs", {}).get("color", "")
            if color_code:
                r, g, b = hex_to_rgb(color_code)
                run.font.color.rgb = RGBColor(r, g, b)
        elif mtype == "fontsize":
            size = mark.get("attrs", {}).get("size", "")
            if size:
                run.font.size = Pt(size)


# --------------------------------------------------------
# 🧱 Node Renderer
# --------------------------------------------------------

def render_node(node, document):
    """Render a node (paragraph, heading, list, etc.) into the document."""
    ntype = node.get("type", "")

    if ntype == "paragraph":
        p = document.add_paragraph()
        for child in node.get("content", []):
            if child.get("type") == "text":
                text_value = sanitize_for_xml(child.get("text", ""))
                run = p.add_run(text_value)
                if "marks" in child:
                    apply_marks(run, child["marks"])

    elif ntype == "bullet_list":
        for item in node.get("content", []):
            p = document.add_paragraph(style='List Bullet')
            for child in item.get("content", []):
                if child.get("type") == "paragraph":
                    for c in child.get("content", []):
                        if c.get("type") == "text":
                            text_value = sanitize_for_xml(c.get("text", ""))
                            run = p.add_run(text_value)
                            if "marks" in c:
                                apply_marks(run, c["marks"])

    elif ntype == "ordered_list":
        for item in node.get("content", []):
            p = document.add_paragraph(style='List Number')
            for child in item.get("content", []):
                if child.get("type") == "paragraph":
                    for c in child.get("content", []):
                        if c.get("type") == "text":
                            text_value = sanitize_for_xml(c.get("text", ""))
                            run = p.add_run(text_value)
                            if "marks" in c:
                                apply_marks(run, c["marks"])

    elif ntype == "heading":
        level = node.get("attrs", {}).get("level", 1)
        p = document.add_paragraph(style=f'Heading {min(level, 9)}')
        for child in node.get("content", []):
            if child.get("type") == "text":
                run = p.add_run(sanitize_for_xml(child.get("text", "")))
                if "marks" in child:
                    apply_marks(run, child["marks"])

    else:
        # fallback: join any loose text content
        text_content = []
        for child in node.get("content", []):
            if isinstance(child, dict) and "text" in child:
                text_content.append(sanitize_for_xml(child["text"]))
        if text_content:
            document.add_paragraph(" ".join(text_content))


# --------------------------------------------------------
# 📄 Document Creation
# --------------------------------------------------------

def render_doc(data):
    """Render an entire document from structured data."""
    doc = Document()
    for node in data.get("content", []):
        render_node(node, doc)
    return doc


def clean_subdoc(sub_doc: Document):
    """Remove escaped WordML text from all paragraphs in a sub-document."""
    for paragraph in sub_doc.paragraphs:
        for run in paragraph.runs:
            run.text = strip_word_xml(run.text)


def append_subdoc(main_doc: Document, sub_doc: Document):
    """Safely append a sub-document into a main document."""
    clean_subdoc(sub_doc)
    for element in sub_doc.element.body:
        main_doc.element.body.append(element)


def converter(filename, data):
    """
    Generate a DOCX file from main_data and optional subdocuments.
    """
    main_doc = render_doc(data)
    main_doc.save(filename)
