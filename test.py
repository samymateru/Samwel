from typing import List
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Font settings
FONT_NAME = "Arial Narrow"
FONT_SIZE = 12

# Map effectiveness to background color (cell background)
effectiveness_bg_colors = {
    "Effective": "00FF00",           # Green
    "Partially Effective": "FFFF00", # Yellow
    "Ineffective": "FF0000",         # Red
}


def set_cell_background(cell, color):
    """Helper to set background color for a Word table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def process_summary_page(programs: List, filename_or_buffer):
    """
    Generate a professional process summary page:
    - Single table
    - Each program gets a full-width header row
    - Sub-programs listed directly below
    - Arial Narrow, size 12 font
    """

    headers = ["Process", "Effectiveness", "Total Findings"]
    document = Document()

    # Create a single table for all programs + subprograms
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        run = hdr_cells[i].paragraphs[0].add_run(header)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)

    # Loop through programs and subprograms
    for idx, program in enumerate(programs, start=1):
        # === Program header row (merged full-width, taller, vertically centered) ===
        program_row = table.add_row()
        program_cell = program_row.cells[0]
        program_cell.merge(program_row.cells[-1])

        # Set a taller row height
        tr = program_row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), "400")  # approx 0.35 inch height
        trHeight.set(qn("w:hRule"), "atLeast")  # ensures minimum height
        trPr.append(trHeight)

        # Vertically center text
        program_cell.vertical_alignment = 1  # 0=top, 1=center, 2=bottom

        # Add text
        p = program_cell.paragraphs[0]
        run = p.add_run(f"{idx}. {program['program']}")
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE + 1)  # Slightly larger font for section header
        p.paragraph_format.space_after = Pt(6)

        sub_programs = program.get("sub_programs", [])
        if not sub_programs:
            empty_row = table.add_row()
            cell = empty_row.cells[0]
            cell.merge(empty_row.cells[-1])
            p = cell.paragraphs[0]
            run = p.add_run("No sub-programs available.")
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)
            p.paragraph_format.left_indent = Inches(0.3)
            continue

        # === Add each sub-program row (with bullet) ===
        for sub in sub_programs:
            row_cells = table.add_row().cells

            title = sub.get("title", "")
            effectiveness = sub.get("effectiveness", "")
            total = str(sub.get("issue_counts", {}).get("total", 0))

            # Add bullet before title
            p = row_cells[0].paragraphs[0]
            run = p.add_run(f"• {title}")
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)

            # Fill the other cells
            row_cells[1].text = effectiveness
            row_cells[2].text = total

            # Apply font style to the rest
            for i in [1, 2]:
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(FONT_SIZE)

            # Set background color for Effectiveness column
            if effectiveness in effectiveness_bg_colors:
                set_cell_background(row_cells[1], effectiveness_bg_colors[effectiveness])

    # Optional space after the table
    document.add_paragraph()

    # Save to file or memory buffer
    if isinstance(filename_or_buffer, BytesIO):
        document.save(filename_or_buffer)
        filename_or_buffer.seek(0)
    else:
        document.save(filename_or_buffer)