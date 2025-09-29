import warnings

from docx import Document

warnings.filterwarnings("ignore", category=UserWarning)
from docxtpl import DocxTemplate, Subdoc



doc = DocxTemplate("name.docx")

# Create a context dictionary with keys matching your placeholders
context = {
    'name': "Samwel",
    'people': [
        {'name': 'Samuel', 'place': 'Python Land'},
        {'name': 'Anna', 'place': 'Data City'},
        {'name': 'Mike', 'place': 'Code Town'},
        {'name': 'Samuel', 'place': 'Python Land'},
    ]
}

# Render the context into the template
doc.render(context)

# Save the result
doc.save("rendered.docx")




def replace_placeholder_with_doc(source_docx_path, insert_docx_path, placeholder):
    # Load the rendered doc with placeholder
    main_doc = Document(source_docx_path)

    # Load the doc you want to insert
    insert_doc = Document(insert_docx_path)

    for para in main_doc.paragraphs:
        if placeholder in para.text:
            # Save style to reuse
            style = para.style

            # Remove the placeholder paragraph
            p_element = para._element
            p_element.getparent().remove(p_element)

            # Insert content
            for insert_para in insert_doc.paragraphs:
                new_para = main_doc.add_paragraph(insert_para.text, style=insert_para.style)
                new_para.style = style
            break


    main_doc.save("final_combined.docx")


replace_placeholder_with_doc("rendered.docx", "placeholder.docx", "[INSERT_CONTENT_HERE]")
